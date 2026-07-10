"""Text extraction for uploaded documents (PDF, DOCX, TXT, MD)."""
import io
from pathlib import Path

MAX_CHARS = 120_000  # safety cap before sending anything to the LLM


def extract_text_from_upload(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Use PDF, DOCX, TXT, or MD.")

    text = _normalize(text)
    if not text.strip():
        raise ValueError(
            "No readable text was found in the file. "
            "If this is a scanned PDF, export a text-based copy and try again."
        )
    return text[:MAX_CHARS]


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _normalize(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned: list[str] = []
    blank = 0
    for line in lines:
        if line.strip():
            blank = 0
            cleaned.append(line)
        else:
            blank += 1
            if blank <= 1:
                cleaned.append("")
    return "\n".join(cleaned)
